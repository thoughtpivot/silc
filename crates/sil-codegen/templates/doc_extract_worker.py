#!/usr/bin/env python3
"""Silc document extract worker (ADR-011).

Reads a staged upload path and emits structured JSON fields:
  title, headings, body, tables, filename, mime, format, char_count

No Pandoc — format-specific Python libs only.
"""

from __future__ import annotations

import argparse
import json
import mimetypes
import os
import re
import sys
from pathlib import Path


SUPPORTED = {".pdf", ".docx", ".odt", ".md", ".markdown", ".txt", ".html", ".htm"}


def guess_format(path: Path) -> str:
    ext = path.suffix.lower()
    return {
        ".pdf": "pdf",
        ".docx": "docx",
        ".odt": "odt",
        ".md": "markdown",
        ".markdown": "markdown",
        ".txt": "text",
        ".html": "html",
        ".htm": "html",
    }.get(ext, ext.lstrip(".") or "unknown")


def extract_txt(path: Path) -> dict:
    text = path.read_text(encoding="utf-8", errors="replace")
    lines = [ln.strip() for ln in text.splitlines()]
    title = next((ln for ln in lines if ln), path.stem)
    return {"title": title, "headings": "", "body": text, "tables": ""}


def extract_md(path: Path) -> dict:
    text = path.read_text(encoding="utf-8", errors="replace")
    headings = []
    for ln in text.splitlines():
        m = re.match(r"^(#{1,6})\s+(.+)$", ln.strip())
        if m:
            headings.append(m.group(2).strip())
    title = headings[0] if headings else path.stem
    return {
        "title": title,
        "headings": "\n".join(headings),
        "body": text,
        "tables": "",
    }


def extract_html(path: Path) -> dict:
    from bs4 import BeautifulSoup

    raw = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    title_el = soup.find("title")
    h_tags = soup.find_all(re.compile(r"^h[1-6]$", re.I))
    headings = [h.get_text(" ", strip=True) for h in h_tags if h.get_text(strip=True)]
    title = (
        (title_el.get_text(" ", strip=True) if title_el else "")
        or (headings[0] if headings else path.stem)
    )
    tables = []
    for table in soup.find_all("table"):
        rows = []
        for tr in table.find_all("tr"):
            cells = [c.get_text(" ", strip=True) for c in tr.find_all(["th", "td"])]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            tables.append("\n".join(rows))
    body = soup.get_text("\n", strip=True)
    return {
        "title": title,
        "headings": "\n".join(headings),
        "body": body,
        "tables": "\n\n".join(tables),
    }


def extract_docx(path: Path) -> dict:
    from docx import Document

    doc = Document(str(path))
    headings = []
    paragraphs = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style = (p.style.name or "") if p.style is not None else ""
        if style.lower().startswith("heading"):
            headings.append(text)
        paragraphs.append(text)
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            tables.append("\n".join(rows))
    title = headings[0] if headings else (paragraphs[0] if paragraphs else path.stem)
    return {
        "title": title,
        "headings": "\n".join(headings),
        "body": "\n".join(paragraphs),
        "tables": "\n\n".join(tables),
    }


def extract_odt(path: Path) -> dict:
    from odf import text as odf_text
    from odf.opendocument import load
    from odf.table import Table, TableCell, TableRow

    doc = load(str(path))
    paragraphs = []
    headings = []
    for el in doc.getElementsByType(odf_text.H):
        t = "".join(n.data for n in el.childNodes if hasattr(n, "data")).strip()
        if t:
            headings.append(t)
            paragraphs.append(t)
    for el in doc.getElementsByType(odf_text.P):
        t = "".join(n.data for n in el.childNodes if hasattr(n, "data")).strip()
        if t:
            paragraphs.append(t)
    tables = []
    for table in doc.getElementsByType(Table):
        rows = []
        for row in table.getElementsByType(TableRow):
            cells = []
            for cell in row.getElementsByType(TableCell):
                cell_text = []
                for p in cell.getElementsByType(odf_text.P):
                    cell_text.append(
                        "".join(n.data for n in p.childNodes if hasattr(n, "data")).strip()
                    )
                cells.append(" ".join(x for x in cell_text if x))
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            tables.append("\n".join(rows))
    title = headings[0] if headings else (paragraphs[0] if paragraphs else path.stem)
    return {
        "title": title,
        "headings": "\n".join(headings),
        "body": "\n".join(paragraphs),
        "tables": "\n\n".join(tables),
    }


def extract_pdf(path: Path) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    body = "\n".join(pages).strip()
    lines = [ln.strip() for ln in body.splitlines() if ln.strip()]
    title = lines[0][:200] if lines else path.stem
    # PDF structure is weak; headings stay empty unless lines look like titles.
    headings = [ln for ln in lines[:20] if len(ln) < 80 and ln == ln.title()]
    return {
        "title": title,
        "headings": "\n".join(headings[:8]),
        "body": body,
        "tables": "",
    }


EXTRACTORS = {
    "pdf": extract_pdf,
    "docx": extract_docx,
    "odt": extract_odt,
    "markdown": extract_md,
    "text": extract_txt,
    "html": extract_html,
}


def extract(path: Path, filename: str | None = None) -> dict:
    if not path.is_file():
        raise FileNotFoundError(f"upload not found: {path}")
    ext = path.suffix.lower()
    if ext not in SUPPORTED:
        raise ValueError(
            f"unsupported document type `{ext or '(none)'}`; "
            f"supported: {', '.join(sorted(SUPPORTED))}"
        )
    fmt = guess_format(path)
    extractor = EXTRACTORS[fmt]
    fields = extractor(path)
    name = filename or path.name
    mime, _ = mimetypes.guess_type(name)
    body = fields.get("body") or ""
    return {
        "title": fields.get("title") or path.stem,
        "headings": fields.get("headings") or "",
        "body": body,
        "tables": fields.get("tables") or "",
        "filename": name,
        "mime": mime or "application/octet-stream",
        "format": fmt,
        "char_count": str(len(body)),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Silc doc::extract worker")
    parser.add_argument("--path", required=True, help="staged upload path")
    parser.add_argument("--filename", default="", help="original filename")
    parser.add_argument("--json", action="store_true", help="emit JSON on stdout")
    args = parser.parse_args()
    path = Path(args.path)
    try:
        result = extract(path, args.filename or None)
    except Exception as exc:  # noqa: BLE001 — worker boundary
        payload = {"ok": False, "error": str(exc)}
        print(json.dumps(payload), flush=True)
        return 1
    if args.json:
        print(json.dumps({"ok": True, **result}), flush=True)
    else:
        print(result.get("body", ""), flush=True)
    return 0


if __name__ == "__main__":
    sys.exit(main())
